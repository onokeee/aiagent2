"""Wordレポートの生成。そのまま配布・回覧できる体裁で作る。

作りの方針:
  - 表紙 → 目次 → 要約 → 本編 → 結論 → 付録 の順。報告書の型に合わせる。
  - 図と表には通し番号とキャプションを付ける（「図3のとおり」と本文から呼べる）。
  - 日本語フォントを明示的に当てる。指定しないと英字フォントが当たり、
    開いた瞬間に体裁が崩れて見える。
  - グラフは画像として貼る（Wordにネイティブのグラフが無いため）。
    画像化できない環境では、同じ内容の表に自動で置き換える。

1セクション = 1つの dict:
    {heading, body, bullets, table:{columns,rows}, image:bytes, caption,
     note, callout, page_break}
"""
from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import config

NAVY = RGBColor(0x1F, 0x3B, 0x5C)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
HILITE = RGBColor(0xC5, 0x5A, 0x11)
INK = RGBColor(0x22, 0x26, 0x2B)
MUTED = RGBColor(0x6B, 0x72, 0x80)
BAND = "F5F8FC"
HEADER_BG = "1F3B5C"
CALLOUT_BG = "FDF3E7"

MAX_TABLE_ROWS = 40


class ReportError(Exception):
    """レポートを作れない理由（そのまま画面に出す）。"""


# =============================================================================
# 体裁の下ごしらえ
# =============================================================================

def _jp_font(run, size=None, bold=None, color=None, name=None):
    """日本語フォントを当てる（東アジア用は XML で直接指定する）。"""
    font = name or config.REPORT_FONT_JA
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _shade(cell_or_par, hex_color: str):
    el = cell_or_par._tc if hasattr(cell_or_par, "_tc") else cell_or_par._p
    pr = el.get_or_add_tcPr() if hasattr(el, "get_or_add_tcPr") else \
        el.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    pr.append(shd)


def _border(par, *, size=6, color="C55A11", where="left"):
    pPr = par._p.get_or_add_pPr()
    borders = pPr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        pPr.append(borders)
    b = OxmlElement(f"w:{where}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size * 4))
    b.set(qn("w:space"), "8")
    b.set(qn("w:color"), color)
    borders.append(b)


def _field(par, instr: str):
    """Wordのフィールド（ページ番号や目次）を入れる。開いたときに計算される。"""
    r1 = par.add_run()._element
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r1.append(fld)
    r2 = par.add_run()._element
    txt = OxmlElement("w:instrText")
    txt.set(qn("xml:space"), "preserve")
    txt.text = instr
    r2.append(txt)
    r3 = par.add_run()._element
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3.append(sep)
    par.add_run("　")                        # 未計算のときに出る仮の文字
    r5 = par.add_run()._element
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r5.append(end)


def _setup_styles(doc):
    """標準スタイルを日本語向けに整える。"""
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), config.REPORT_FONT_JA)
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before in (("Heading 1", 16, NAVY, 18),
                                      ("Heading 2", 13, NAVY, 14),
                                      ("Heading 3", 11.5, ACCENT, 10)):
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        rPr = st.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), config.REPORT_FONT_JA)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


def _setup_page(doc, footer_text: str):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_text:
        r = p.add_run(footer_text + "　　")
        _jp_font(r, size=8.5, color=MUTED)
    _field(p, "PAGE")
    r = p.add_run(" / ")
    _jp_font(r, size=8.5, color=MUTED)
    _field(p, "NUMPAGES")
    for r in p.runs:
        _jp_font(r, size=8.5, color=MUTED)


def _para(doc, text="", *, size=10.5, bold=False, color=INK, align=None,
          space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        _jp_font(r, size=size, bold=bold, color=color)
    return p


def _fmt(v) -> str:
    if v is None:
        return ""
    if isinstance(v, bool):
        return "はい" if v else "いいえ"
    if isinstance(v, float):
        if v == int(v) and abs(v) < 1e15:
            return f"{int(v):,}"
        return f"{v:,.2f}".rstrip("0").rstrip(".")
    if isinstance(v, int):
        return f"{v:,}"
    return str(v)


# =============================================================================
# 部品
# =============================================================================

def _cover(doc, args: dict):
    for _ in range(4):
        doc.add_paragraph()
    _para(doc, args.get("title", "レポート"), size=26, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    if args.get("subtitle"):
        _para(doc, args["subtitle"], size=13, color=MUTED,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)

    # 表紙の線
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _border(line, size=3, color="1F3B5C", where="bottom")

    for _ in range(6):
        doc.add_paragraph()
    org = args.get("org") or config.REPORT_ORG
    for text in (args.get("date") or datetime.now().strftime("%Y年%m月%d日"),
                 org, args.get("author")):
        if text:
            _para(doc, text, size=11, color=MUTED,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    doc.add_page_break()


def _toc(doc):
    _para(doc, "目次", size=15, bold=True, color=NAVY, space_after=10)
    p = doc.add_paragraph()
    _field(p, r'TOC \o "1-2" \h \z \u')
    _para(doc, "※ 目次はWordで開いたあと、この部分を選んで F9 を押すと最新になります。",
          size=8.5, color=MUTED, space_after=0)
    doc.add_page_break()


def _summary_box(doc, points: list):
    if not points:
        return
    _para(doc, "要点", size=12, bold=True, color=HILITE, space_after=4)
    for s in points:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(str(s))
        _jp_font(r, size=11, bold=True, color=INK)
    doc.add_paragraph()


def _callout(doc, text: str, label="ポイント"):
    p = _para(doc, f"【{label}】{text}", size=10.5, color=INK, space_after=10)
    _border(p, size=6, color="C55A11", where="left")
    _shade(p, CALLOUT_BG)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)


def _table(doc, columns, rows, *, caption=None, number=None, note=None):
    limit = MAX_TABLE_ROWS
    shown, cut = rows[:limit], max(0, len(rows) - limit)
    t = doc.add_table(rows=len(shown) + 1, cols=len(columns))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    for j, c in enumerate(columns):
        cell = t.cell(0, j)
        cell.text = ""
        _shade(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(str(c))
        _jp_font(r, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(shown, start=1):
        for j in range(len(columns)):
            v = row[j] if j < len(row) else ""
            cell = t.cell(i, j)
            cell.text = ""
            if i % 2 == 0:
                _shade(cell, BAND)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            if isinstance(v, (int, float)) and not isinstance(v, bool):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(_fmt(v))
            _jp_font(r, size=9.5)

    tail = []
    if caption:
        cap = _para(doc, f"表{number} {caption}" if number else caption,
                    size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_after=4)
        cap.paragraph_format.space_before = Pt(3)
    if cut:
        tail.append(f"全 {len(rows):,} 行のうち上位 {limit} 行を掲載")
    if note:
        tail.append(note)
    if tail:
        _para(doc, "　".join(tail), size=8.5, color=MUTED, space_after=10)
    return t


def _image(doc, data: bytes, *, caption=None, number=None, width_cm=16.0):
    doc.add_picture(io.BytesIO(data), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        _para(doc, f"図{number} {caption}" if number else caption, size=9,
              color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)


# =============================================================================
# 組み立て
# =============================================================================

def build(sections: list[dict], *, title="レポート", subtitle="", summary=None,
          conclusion="", recommendations=None, caveats=None, footer="",
          org="", author="", toc=True, appendix=None) -> bytes:
    """セクションのリストから .docx のバイト列を作る。"""
    if not sections:
        raise ReportError("セクションが1つもありません。")

    doc = Document()
    _setup_styles(doc)
    _setup_page(doc, footer or config.REPORT_ORG or "")
    _cover(doc, {"title": title, "subtitle": subtitle, "org": org, "author": author})
    if toc:
        _toc(doc)

    if summary:
        doc.add_heading("要約", level=1)
        _summary_box(doc, summary)

    fig_no, tbl_no = 0, 0
    for i, s in enumerate(sections, 1):
        if not s.get("heading"):
            raise ReportError(f"{i}番目のセクションに heading がありません。")
        if s.get("page_break"):
            doc.add_page_break()
        doc.add_heading(s["heading"], level=int(s.get("level") or 1))
        if s.get("body"):
            _para(doc, s["body"])
        for b in (s.get("bullets") or []):
            text = b.get("text", "") if isinstance(b, dict) else str(b)
            level = int(b.get("level", 0)) if isinstance(b, dict) else 0
            p = doc.add_paragraph(style="List Bullet" if not level
                                  else "List Bullet 2")
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            _jp_font(r, size=10.5)
        if s.get("image"):
            fig_no += 1
            _image(doc, s["image"], caption=s.get("caption") or s["heading"],
                   number=fig_no)
        if s.get("table"):
            t = s["table"]
            if not t.get("columns"):
                raise ReportError(f"「{s['heading']}」の表に columns がありません。")
            tbl_no += 1
            _table(doc, t["columns"], t.get("rows") or [],
                   caption=s.get("table_caption") or s.get("caption") or s["heading"],
                   number=tbl_no, note=t.get("note"))
        if s.get("callout"):
            _callout(doc, s["callout"])
        if s.get("note"):
            p = _para(doc, s["note"], size=10, color=MUTED, space_after=10)
            _border(p, size=4, color="D5DBE2", where="left")
            p.paragraph_format.left_indent = Cm(0.4)

    if conclusion:
        doc.add_heading("結論", level=1)
        _para(doc, conclusion)
    if recommendations:
        doc.add_heading("推奨する打ち手", level=1)
        for i, r in enumerate(recommendations, 1):
            if isinstance(r, dict):
                text = r.get("text", "")
                extra = "　".join(x for x in
                                  (f"担当: {r['owner']}" if r.get("owner") else "",
                                   f"期限: {r['due']}" if r.get("due") else "") if x)
            else:
                text, extra = str(r), ""
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"{i}. ")
            _jp_font(run, size=10.5, bold=True, color=HILITE)
            run = p.add_run(text)
            _jp_font(run, size=10.5, bold=True)
            if extra:
                run = p.add_run(f"（{extra}）")
                _jp_font(run, size=9.5, color=MUTED)
    if caveats:
        doc.add_heading("前提・注意", level=1)
        for c in caveats:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(str(c))
            _jp_font(r, size=10, color=MUTED)

    for extra in (appendix or []):
        doc.add_page_break()
        doc.add_heading(extra.get("heading", "付録"), level=1)
        if extra.get("body"):
            _para(doc, extra["body"])
        if extra.get("table"):
            tbl_no += 1
            _table(doc, extra["table"]["columns"], extra["table"].get("rows") or [],
                   caption=extra.get("caption"), number=tbl_no)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_filename(name: str | None, default: str = "report") -> str:
    base = re.sub(r'[\\/:*?"<>|]', "_", str(name or default)).strip() or default
    return base if base.lower().endswith(".docx") else base + ".docx"


def outline(sections: list[dict]) -> list[str]:
    out = []
    for i, s in enumerate(sections, 1):
        bits = []
        if s.get("image"):
            bits.append("図")
        if s.get("table"):
            bits.append(f"表{len(s['table'].get('rows') or [])}行")
        if s.get("bullets"):
            bits.append(f"箇条書き{len(s['bullets'])}件")
        out.append(f"{i}. {s.get('heading', '')}"
                   + (f"（{'・'.join(bits)}）" if bits else ""))
    return out
